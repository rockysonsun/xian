#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
创建符合GB/T 9704-2012标准的政府公文格式DOCX
"""

from docx import Document
from docx.shared import Pt, Mm, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            for key in ["sz", "val", "color", "space"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))
            tcPr.append(element)

def set_run_font(run, font_name='仿宋', font_size=16, bold=False):
    """设置字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)

def add_red_line(doc):
    """添加红色分隔线（模拟）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('─' * 50)
    run.font.color.rgb = RGBColor(255, 0, 0)
    run.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def create_gov_document():
    doc = Document()
    
    # ===== 页面设置 =====
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    # 公文标准：上37mm，下35mm，左28mm，右26mm
    section.top_margin = Mm(37)
    section.bottom_margin = Mm(35)
    section.left_margin = Mm(28)
    section.right_margin = Mm(26)
    
    # ===== 版头（红色） =====
    # 发文机关标志
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('××单位文件')
    run.font.name = '方正小标宋简体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(255, 0, 0)
    run.font.bold = False
    
    # 发文字号
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('××发〔2026〕××号')
    set_run_font(run, '仿宋', 14)
    
    # 红色分隔线
    add_red_line(doc)
    
    # ===== 主体 =====
    
    # 标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(18)
    p.paragraph_format.line_spacing = Pt(36)
    run = p.add_run('关于推动光交换技术（OCS）产业发展的报告')
    run.font.name = '方正小标宋简体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')
    run.font.size = Pt(22)
    run.font.bold = False
    
    # 主送机关
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run('××××领导：')
    set_run_font(run, '仿宋', 16)
    
    # ===== 正文内容 =====
    
    # 导语
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)  # 2字符
    p.paragraph_format.line_spacing = Pt(29)  # 固定值29磅
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('随着人工智能技术的快速发展，算力已成为数字经济时代的核心生产力。光交换技术（OCS）作为算力基础设施的关键技术，对于提升我国算力水平、保障算力安全自主可控具有重要意义。现将有关情况报告如下：')
    set_run_font(run, '仿宋', 16)
    
    # 第一部分标题（黑体）
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(29)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('一、光交换技术的基本概念')
    set_run_font(run, '黑体', 16, bold=True)
    
    # 正文
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(29)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('光交换技术（Optical Circuit Switching，简称OCS）是一种利用光信号进行数据交换的新型网络技术。与传统的电交换技术相比，光交换技术具有传输速度快、能耗低、延迟小等显著优势。')
    set_run_font(run, '仿宋', 16)
    
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(29)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('通俗地说，如果把数据比作车辆，网络比作道路，那么电交换技术就像传统的收费站，每辆车都需要停车缴费后才能通行；而光交换技术则像高速公路的ETC通道，车辆可以不停车快速通过，大大提高了通行效率。')
    set_run_font(run, '仿宋', 16)
    
    # 第二部分
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(29)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('二、发展光交换技术的重要意义')
    set_run_font(run, '黑体', 16, bold=True)
    
    # （一）小标题（楷体）
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(29)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('（一）保障国家算力安全')
    set_run_font(run, '楷体', 16, bold=True)
    
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(29)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('当前，我国高端网络设备核心芯片严重依赖进口，存在被"卡脖子"风险。光交换技术基于硅光工艺，不依赖最先进的芯片制程，可以实现完全自主可控，是保障国家算力安全的重要突破口。')
    set_run_font(run, '仿宋', 16)
    
    # （二）
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(29)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('（二）提升AI算力集群效率')
    set_run_font(run, '楷体', 16, bold=True)
    
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(29)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('当前，人工智能大模型训练需要数千甚至上万台服务器协同工作，这些服务器之间每秒要交换数百万亿次数据。传统电交换技术就像"收费站"，每批数据都需要停车处理，已成为算力集群的瓶颈。光交换技术（OCS）能够在机房内部实现服务器之间的"光路直达"，无需反复光电转换，使AI训练效率提升约40%，是建设高效智算中心的关键技术。')
    set_run_font(run, '仿宋', 16)
    
    # （三）
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(29)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('（三）实现绿色低碳发展')
    set_run_font(run, '楷体', 16, bold=True)
    
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(29)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('数据中心是能耗大户，其中网络设备耗电占比高达30%以上。光交换技术相比传统电交换技术可节能40%至70%，且几乎不发热，大幅减少空调用电。据测算，一个万卡级智算中心采用光交换技术后，每年可节省电费数千万元，助力实现"双碳"目标。')
    set_run_font(run, '仿宋', 16)
    
    # 第三部分
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(29)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('三、我国光交换技术发展现状')
    set_run_font(run, '黑体', 16, bold=True)
    
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(29)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('我国光交换技术发展迅速，已具备较好的产业基础。光迅科技等企业已实现192×192端口光交换芯片量产，320×320端口芯片完成演示，技术水平与国际基本同步。同时，我国是全球除美国外唯一具备完整光交换产业链的国家，从芯片设计、设备制造到系统集成均可自主完成。')
    set_run_font(run, '仿宋', 16)
    
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(29)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('2026年4月，工业和信息化部发布《关于开展普惠算力赋能中小企业发展专项行动的通知》，明确提出"推动全光交换等技术应用部署，降低算力应用终端到服务器的网络时延"，首次将光交换技术列为算力网络关键技术路线，为产业发展提供了强有力的政策支持。')
    set_run_font(run, '仿宋', 16)
    
    # 第四部分
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(29)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('四、面临的挑战')
    set_run_font(run, '黑体', 16, bold=True)
    
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(29)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('一是高端芯片仍有差距。大端口光交换芯片（256×256以上）仍处于研发阶段，与国外先进水平存在2至3年差距。二是标准话语权不足。国际标准组织中的光交换标准制定主要由国外企业主导，我国参与度较低。三是规模化应用刚刚起步。国外谷歌公司已部署8年，我国尚处于示范验证阶段。')
    set_run_font(run, '仿宋', 16)
    
    # 第五部分
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(29)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('五、相关建议')
    set_run_font(run, '黑体', 16, bold=True)
    
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(29)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('（一）加强顶层设计。建议将光交换技术列入"十五五"国家信息化规划重大技术方向，设立专项基金支持产业发展。')
    set_run_font(run, '仿宋', 16)
    
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(29)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('（二）推动产业协同。成立国家光交换创新联合体，整合芯片设计、设备制造、系统集成等产业链资源，避免重复投入、各自为战。')
    set_run_font(run, '仿宋', 16)
    
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(29)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('（三）建设示范工程。在新建智算中心中开展光交换技术应用示范，优先在AI训练集群中部署，验证技术可行性，形成可复制推广的建设经验。')
    set_run_font(run, '仿宋', 16)
    
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(29)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('（四）参与标准制定。支持我国原创技术"拓扑感知"上升为国家标准，并推动成为国际标准，掌握技术话语权。')
    set_run_font(run, '仿宋', 16)
    
    # 结语
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(29)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('光交换技术是智算中心内部网络的关键底座，关系到国家算力安全、AI产业发展和绿色低碳目标实现。当前全球光交换技术正处于从技术验证向规模化商用过渡的关键期，2026年至2028年是技术路线定型的窗口期。建议相关部门高度重视，抓住窗口期，统筹规划、加大投入、协同推进，推动我国光交换产业实现从"跟跑"到"并跑"再到"领跑"的跨越。')
    set_run_font(run, '仿宋', 16)
    
    # 结束语
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(29)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('特此报告。')
    set_run_font(run, '仿宋', 16)
    
    # ===== 版记 =====
    # 空行分隔
    for _ in range(3):
        doc.add_paragraph()
    
    # 分隔线
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('─' * 50)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 抄送
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('抄送：××××，××××。')
    set_run_font(run, '仿宋', 14)
    
    # 印发机关和日期
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run1 = p.add_run('××单位办公室')
    set_run_font(run1, '仿宋', 14)
    
    # 右对齐日期
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('2026年5月9日印发')
    set_run_font(run, '仿宋', 14)
    
    return doc

if __name__ == '__main__':
    doc = create_gov_document()
    output_path = '/Users/rocky/.openclaw/workspace/reports/光交换技术发展报告_政府公文版.docx'
    doc.save(output_path)
    print(f"文档已保存至: {output_path}")
