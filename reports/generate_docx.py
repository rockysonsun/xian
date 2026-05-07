#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成政府标准格式的OCS建议书docx文件
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def set_run_font(run, font_name='仿宋', size=16, bold=False):
    """设置字体"""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_heading_custom(doc, text, level=1):
    """添加自定义标题"""
    if level == 0:  # 文档标题
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, '方正小标宋简体', 22, False)
        p.paragraph_format.space_after = Pt(30)
        p.paragraph_format.space_before = Pt(0)
    elif level == 1:  # 第一章
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_run_font(run, '黑体', 16, True)
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(12)
    elif level == 2:  # 一、
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_run_font(run, '楷体_GB2312', 16, True)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(10)
    elif level == 3:  # （一）
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_run_font(run, '仿宋_GB2312', 16, True)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(8)
    elif level == 4:  # 1.
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_run_font(run, '仿宋_GB2312', 16, True)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    return p

def add_body_text(doc, text, indent=True, bold=False):
    """添加正文"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)  # 2字符
    run = p.add_run(text)
    set_run_font(run, '仿宋_GB2312', 16, bold)
    p.paragraph_format.line_spacing = 1.8
    p.paragraph_format.space_after = Pt(6)
    return p

def add_highlight_text(doc, text):
    """添加高亮文字（加粗）"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, '仿宋_GB2312', 16, True)
    p.paragraph_format.line_spacing = 1.8
    return p

def add_quote(doc, text):
    """添加引用框"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    run = p.add_run(text)
    set_run_font(run, '仿宋_GB2312', 16, False)
    run.italic = True
    p.paragraph_format.line_spacing = 1.8
    # 添加左边框效果
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '12')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), '666666')
    pBdr.append(left)
    pPr.append(pBdr)
    return p

def add_table_custom(doc, headers, rows, caption=None):
    """添加表格（三线表）"""
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(caption)
        set_run_font(run, '黑体', 14, False)
        p.paragraph_format.space_after = Pt(6)
    
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # 设置表头
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_run_font(run, '黑体', 14, False)
        set_cell_shading(hdr_cells[i], 'F5F5F5')
    
    # 设置数据行
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = str(cell_text)
            for paragraph in row_cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, '仿宋_GB2312', 14, False)
    
    # 设置表格边框为三线表样式
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    
    # 添加表格后间距
    doc.add_paragraph()
    
    return table

def main():
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    
    # ========== 封面 ==========
    add_heading_custom(doc, '中国OCS（光电路交换）发展路线建议书', 0)
    
    # 元信息
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    meta_text = '编制单位：星宿老仙技术团队\n编制日期：2026年4月13日\n版本：V1.2.2（已整合工信部政策文件，全面修正OCS产业数据）\n密级：内部参考'
    run = meta.add_run(meta_text)
    set_run_font(run, '仿宋_GB2312', 14, False)
    meta.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()  # 空行
    
    # ========== 摘要 ==========
    add_heading_custom(doc, '摘要', 2)
    
    add_body_text(doc, '本建议书基于对全球OCS技术发展趋势、中国产业现状、政策环境以及"东数西算"等国家战略的深入分析，并充分整合SuperPod白皮书《分布式AI推理加速架构技术白皮书》的核心技术框架（包括系统能力边界外移理论、帕累托前沿分析方法论、三大OCS探索构型），提出适合中国国情的OCS发展路线建议。')
    
    add_highlight_text(doc, '建议书核心主张：')
    add_quote(doc, '以"拓扑感知"等原创技术为突破口，通过"标准引领+产业协同+场景驱动"三位一体策略，在2026-2030年实现中国OCS产业从跟跑到并跑、部分领域领跑的跨越。')
    
    add_highlight_text(doc, '政策背景（2026年4月最新）：')
    add_quote(doc, '2026年4月2日，工信部办公厅发布《关于开展普惠算力赋能中小企业发展专项行动的通知》，明确提出"推动全光交换等技术应用部署，降低算力应用终端到服务器的网络时延，提升应用交互体验"。这标志着OCS技术已从产业探索上升为国家政策导向，为产业发展提供了强有力的政策支撑。')
    
    add_highlight_text(doc, '核心判断（基于SuperPod白皮书）：')
    add_quote(doc, '超节点竞争的本质不是比拼任何单一规格，而是比拼系统能力边界外移的速度。系统能力边界外移速度的差异，会以指数形式放大为系统性能差距。四代之后差距即达16倍。')
    
    add_body_text(doc, '这意味着OCS不仅是一项技术选择，而是决定中国智算产业能否避免被拉开"难以逆转的差距"的关键变量。本建议书提出的发展路线，正是基于这一核心判断，为中国OCS产业提供从架构到落地的完整技术路径。')
    
    # ========== 第一章 ==========
    add_heading_custom(doc, '第一章 形势分析：为什么中国必须发展OCS', 1)
    
    add_heading_custom(doc, '一、全球OCS发展态势', 2)
    
    add_heading_custom(doc, '（一）技术突破期（2026年为关键拐点）', 3)
    
    add_table_custom(doc, 
        ['维度', '全球现状', '中国现状', '差距评估'],
        [
            ['技术成熟度', '谷歌TPU v4-v7已规模化部署8年', '光迅192×192量产，320×320演示', '基本同步'],
            ['产业链完整性', '美、日、欧主导核心器件', '全链条覆盖，部分环节领先', '局部领先'],
            ['市场规模', '2025年7.83亿美元', '2025年约22亿元人民币', '全球第二'],
            ['标准话语权', 'OCP OCS项目组2025年成立', '积极参与，尚未主导', '需突破']
        ],
        '表1-1 全球与中国OCS技术发展对比'
    )
    
    add_heading_custom(doc, '（二）关键驱动因素', 3)
    
    add_body_text(doc, '1. AI算力需求爆发：模型参数每年10倍增长，GPU集群从十万卡向百万卡迈进；')
    add_body_text(doc, '2. 电交换瓶颈凸显：功耗墙（单台51.2T电交换机超3000W）、时延（数十微秒）、距离（小于1.5米）；')
    add_body_text(doc, '3. OCS技术优势：功耗降低40%-70%、时延微秒/纳秒级、协议透明、带宽密度高；')
    add_body_text(doc, '4. 国家政策推动：工信部明确支持全光交换技术应用部署（2026年4月）。')
    
    add_heading_custom(doc, '二、政策环境：国家层面明确支持', 2)
    
    add_heading_custom(doc, '（一）工信部最新政策导向（2026年4月）', 3)
    
    add_highlight_text(doc, '文件依据：')
    add_quote(doc, '2026年4月2日，工业和信息化部办公厅发布《关于开展普惠算力赋能中小企业发展专项行动的通知》（工信厅通信〔2026〕XX号）')
    
    add_highlight_text(doc, '核心内容：')
    add_body_text(doc, '一是增强算网融合支撑能力。深入开展城域"毫秒用算"专项行动，扩大城域1毫秒时延圈覆盖范围。')
    add_body_text(doc, '二是推动全光交换技术应用。明确提出"推动全光交换等技术应用部署，降低算力应用终端到服务器的网络时延，提升应用交互体验"。')
    add_body_text(doc, '三是打通最后一公里。提升网络对算力服务的支撑能力，打通"最后一公里"网络和算力接入瓶颈，降低中小企业算力访问时延。')
    add_body_text(doc, '四是光网络延伸。推动光网络设备向产业园区等用户侧延伸。')
    
    add_highlight_text(doc, '政策意义：')
    add_table_custom(doc,
        ['维度', '政策价值'],
        [
            ['技术路线确认', '国家层面首次明确将全光交换（OCS）作为算力网络关键技术路线'],
            ['产业信心提振', '为OCS产业发展提供政策背书，降低市场不确定性'],
            ['市场需求释放', '运营商、云厂商将加速OCS采购和部署'],
            ['标准制定加速', '为OCS国家标准、行业标准制定提供政策依据']
        ],
        '表1-2 工信部政策对OCS产业的意义'
    )
    
    add_heading_custom(doc, '（二）算力安全自主可控', 3)
    add_body_text(doc, '现状风险：高端电交换芯片（51.2T/102.4T）依赖进口，存在断供风险。')
    add_body_text(doc, 'OCS优势：硅光技术基于CMOS工艺，不依赖先进制程（7nm/5nm），可实现自主可控。')
    add_body_text(doc, '关键数据：中国信通院报告显示，算力每投入1元，带动3-4元经济产出。')
    
    add_heading_custom(doc, '（三）"东数西算"工程需求', 3)
    add_table_custom(doc,
        ['工程目标', 'OCS价值'],
        [
            ['全国一体化算力网', 'OCS实现算力枢纽间全光互联，时延小于10ms'],
            ['绿色低碳（PUE小于1.25）', 'OCS功耗较电交换降低40%-70%'],
            ['毫秒用算', 'OCS微秒级切换，支撑实时算力调度'],
            ['八大枢纽互联', 'OCS简化网络层级，光模块使用量减少50%']
        ],
        '表1-3 "东数西算"工程目标与OCS价值对应'
    )
    
    add_heading_custom(doc, '（四）产业换道超车机遇', 3)
    add_body_text(doc, '传统电交换：思科、Arista等垄断，中国难以突破。')
    add_body_text(doc, 'OCS新赛道：全球处于同一起跑线，中国具备全产业链优势。')
    add_body_text(doc, '窗口期：2026-2028年是技术路线定型关键期，错过将失去话语权。')
    
    # ========== 第二章 ==========
    add_heading_custom(doc, '第二章 中国OCS产业现状评估', 1)
    
    add_heading_custom(doc, '一、产业链全景图', 2)
    add_body_text(doc, '中国OCS产业链已初步形成完整布局，涵盖上游支撑、中游核心、下游应用及标准生态四个环节。上游包括上海微电子90nm光刻、中芯国际8英寸硅光、赛微电子MEMS代工、罗博特科键合设备等；中游包括光迅科技（192×192量产）、德科立（320×320演示中）、华工正源、联特科技、曦智科技、中兴通讯等；下游包括阿里云、腾讯云、字节跳动、华为、壁仞科技及三大运营商；标准生态方面包括OCP中国（OCTC）、中国通信标准化协会、开放数据中心委员会等。')
    
    add_heading_custom(doc, '二、优势与短板分析', 2)
    
    add_heading_custom(doc, '（一）核心优势', 3)
    add_table_custom(doc,
        ['优势领域', '具体表现', '战略价值'],
        [
            ['全产业链覆盖', '全球除美国外唯一具备完整OCS产业链', '供应链安全、成本优势'],
            ['成本竞争力', '量产成本较海外低30%-40%', '规模化替代基础'],
            ['市场需求巨大', '全球最大AI算力、数据中心市场', '应用场景丰富'],
            ['技术追赶快', '光迅192×192量产，320×320演示', '差距持续缩小'],
            ['政策支持力度大', '"东数西算"、"毫秒用算"等专项', '发展环境优越']
        ],
        '表2-1 中国OCS产业核心优势'
    )
    
    add_heading_custom(doc, '（二）关键短板', 3)
    add_table_custom(doc,
        ['短板领域', '具体表现', '风险等级'],
        [
            ['大端口芯片', '海外128×128验证中，国内64×64研发中', '高'],
            ['高端设备依赖', 'DUV光刻机、硅光EDA依赖进口', '高'],
            ['标准话语权', 'OCP OCS项目组刚成立，中国参与度低', '中'],
            ['规模化商用', '海外谷歌已部署8年，国内刚起步', '中'],
            ['高端人才', '硅光设计、系统架构人才短缺', '中']
        ],
        '表2-2 中国OCS产业关键短板'
    )
    
    add_heading_custom(doc, '三、竞争格局分析', 2)
    
    add_heading_custom(doc, '（一）国内主要玩家', 3)
    add_table_custom(doc,
        ['企业', '技术路线', '产品状态', '核心优势', '代表客户'],
        [
            ['光迅科技', 'MEMS', '192×192量产，320×320演示', '全栈自研、谷歌供应链', '谷歌、华为、运营商'],
            ['德科立', '硅基波导', '128×128样机', '微秒级切换、海外客户', '谷歌（送样验证）'],
            ['华工正源', '硅光', '64×64样机（2025 OFC发布）', '光模块协同、云厂商资源', '阿里云'],
            ['曦智科技', '硅光', '128卡超节点商用', '全球首创、超节点方案', '壁仞科技、中兴'],
            ['中兴通讯', '电信级OCS', '整机量产', '运营商渠道、系统集成', '三大运营商']
        ],
        '表2-3 国内OCS主要企业竞争格局'
    )
    
    add_heading_custom(doc, '（二）国际对标', 3)
    add_table_custom(doc,
        ['维度', '谷歌（美国）', '中国头部企业', '差距'],
        [
            ['部署规模', 'TPU集群8年，万台级', '刚起步，百台级', '3-5年'],
            ['端口规模', '136×136（量产），320×320（研发）', '192×192（量产），320×320（演示）', '基本同步'],
            ['技术成熟度', '自研Palomar MEMS，成熟', '多种路线并行，探索中', '2-3年'],
            ['生态掌控', '自研ASIC+OCS垂直整合', '依赖外部GPU/ASIC', '生态短板']
        ],
        '表2-4 中美OCS产业国际对标'
    )
    
    # ========== 第三章 ==========
    add_heading_custom(doc, '第三章 发展路线建议', 1)
    
    add_heading_custom(doc, '一、总体思路', 2)
    add_highlight_text(doc, '"三步走"战略：')
    add_body_text(doc, '第一步（2026-2027）：重点突破，局部领先。实现32×64端口硅光芯片量产，拓扑感知等原创技术标准化，东数西算枢纽节点试点部署。')
    add_body_text(doc, '第二步（2028-2029）：规模商用，并跑发展。实现64×64/128×128高端口芯片突破，国产AI芯片+OCS协同方案成熟，云厂商、运营商规模采购。')
    add_body_text(doc, '第三步（2030+）：生态引领，部分领跑。实现256×256+超大端口芯片，OCS+CPO全光互联架构普及，中国标准走向国际。')
    
    add_heading_custom(doc, '二、技术路线建议', 2)
    
    add_heading_custom(doc, '（一）双轮驱动：MEMS+硅光并行', 3)
    add_table_custom(doc,
        ['技术路线', '适用场景', '发展重点', '目标企业'],
        [
            ['MEMS OCS', '大端口、低频次切换', '320×320+超大端口，小于1.5dB插入损耗', '光迅科技、赛微电子'],
            ['硅光OCS', '中小端口、高频次切换', '纳秒级切换，与CPO集成', '德科立、曦智科技、华工正源']
        ],
        '表3-1 MEMS与硅光技术路线对比'
    )
    
    add_heading_custom(doc, '（二）三大探索构型：适配不同场景需求', 3)
    add_body_text(doc, '根据SuperPod白皮书的技术分析，结合中国产业现状，建议重点布局以下三种OCS架构：')
    
    add_heading_custom(doc, '1. 构型一：Dragonfly+OCS（规模优先型）', 4)
    add_body_text(doc, '架构特点：组内电交换（Clos）+组间光交换（OCS），适合万卡级超节点，规模扩展性最强，复用以太网生态（RoCEv2/UEC），软件适配成本低。代表：英伟达Spectrum-X。')
    add_highlight_text(doc, '发展建议：')
    add_body_text(doc, '近期目标：实现256端口OCS与51.2T电交换机的协同部署。')
    add_body_text(doc, '中期目标：突破512端口OCS，支撑32K GPU规模。')
    add_body_text(doc, '适用场景：通用GPU集群、大规模训练中心。')
    
    add_heading_custom(doc, '2. 构型二：3D Torus+OCS（能效优先型）', 4)
    add_body_text(doc, '架构特点：全光层互联，无大型电交换机，拓扑可动态重构（Ring/Mesh/Torus），网络成本小于5%，网络功耗小于3%（系统占比）。代表：谷歌TPU v4/v7（Palomar OCS）。')
    add_highlight_text(doc, '发展建议：')
    add_body_text(doc, '近期目标：实现64-256节点Cube内电连接+OCS光互联。')
    add_body_text(doc, '中期目标：突破4096-9216卡Pod规模。')
    add_body_text(doc, '适用场景：AI专用集群、高能效智算中心。')
    
    add_heading_custom(doc, '3. 构型三：大环路+dOCS（动态重构型）', 4)
    add_body_text(doc, '架构特点：基于拓扑感知的动态光路调整，支持故障秒级切换、业务无损迁移，与"拓扑感知"专利技术高度契合。代表：华为DC-OXC+拓扑感知专利。')
    add_highlight_text(doc, '发展建议：')
    add_body_text(doc, '近期目标：拓扑感知技术标准化，集成到商用OCS设备。')
    add_body_text(doc, '中期目标：实现毫秒级故障切换、分钟级业务调度。')
    add_body_text(doc, '适用场景：运营商骨干网、高可靠金融算力中心。')
    
    add_heading_custom(doc, '（三）三阶段发展路径（与SuperPod白皮书对齐）', 3)
    add_table_custom(doc,
        ['阶段', '时间', '核心任务', '技术重点', '规模目标'],
        [
            ['近期', '2026-2027', '标准构型落地', 'LPO、HBM4、低精度优化', '8K GPU'],
            ['中期', '2028-2029', '探索构型验证', 'OCS、CPO、Chiplet+UCIe', '32K GPU'],
            ['长期', '2030+', '全栈能力构建', '芯片出光、全光直连', '64K+ GPU']
        ],
        '表3-2 OCS技术三阶段发展路径'
    )
    
    add_heading_custom(doc, '（四）原创技术突破：拓扑感知', 3)
    add_highlight_text(doc, '核心建议：')
    add_body_text(doc, '将"基于两端光模块写码的OCS光传输拓扑快速认知与动态调整方法"（专利号：20260052-P26BJ0080CNCN）上升为国家标准，并推动成为国际标准。')
    add_highlight_text(doc, '技术价值：')
    add_body_text(doc, '解决OCS网络配置复杂痛点；实现拓扑秒级自动感知；提升网络可靠性和运维效率。')
    add_highlight_text(doc, '产业化路径：')
    add_body_text(doc, '2026年：与华为、中兴等设备商合作，集成到DC-OXC/OCS设备。')
    add_body_text(doc, '2027年：在"东数西算"枢纽节点试点验证。')
    add_body_text(doc, '2028年：向OCP提交标准提案，推动国际标准化。')
    
    add_heading_custom(doc, '（五）关键技术攻关清单', 3)
    add_table_custom(doc,
        ['技术方向', '攻关目标', '时间节点', '承担主体'],
        [
            ['高端口硅光芯片', '128×128验证，256×256攻关', '2027年', '光迅、德科立'],
            ['超低损耗MEMS', '插入损耗小于1.0dB，384×384端口', '2028年', '光迅、赛微电子'],
            ['微秒级切换', '切换时延小于1μs', '2028年', '德科立、曦智'],
            ['OCS+CPO集成', '光开关与光引擎单片集成', '2029年', '华工正源、曦智'],
            ['智能调度算法', 'AI驱动的拓扑动态优化', '2027年', '云厂商、算法公司']
        ],
        '表3-3 OCS关键技术攻关清单'
    )
    
    add_heading_custom(doc, '三、产业生态建议', 2)
    
    add_heading_custom(doc, '（一）构建"国家队+民企"协同创新体', 3)
    add_highlight_text(doc, '协同框架（基于SuperPod白皮书产业协同理念）：')
    add_body_text(doc, '国家OCS创新联合体分为三个层面：芯片设计层（光迅科技、德科立、曦智科技、华工正源）、系统设备层（华为、中兴通讯、烽火通信）、应用示范层（阿里云、三大运营商、壁仞科技）。底层支撑共性技术平台包括硅光工艺平台（中芯国际/华虹）、MEMS代工平台（赛微电子）、测试验证中心（信通院）、EDA工具开发（联合攻关）。')
    
    add_quote(doc, '超节点竞争不是单产品竞争，而是跨芯片、互联、光学、封装、整机、软件与运维的协同工程。建立这种协同能力需要产学研用各方的深度参与——芯片厂商提供算力底座、光模块厂商提供互联能力、封装厂商突破集成极限、云厂商验证规模化落地、高校和研究机构提供前沿理论和算法支撑。')
    
    add_highlight_text(doc, '中国OCS产业协同重点：')
    add_table_custom(doc,
        ['参与方', '核心能力', '协同任务'],
        [
            ['芯片厂商（光迅、德科立、曦智）', 'OCS芯片设计', '提供算力底座，突破大端口芯片'],
            ['光模块厂商（华工正源、联特）', '高速光互联', '提供互联能力，协同CPO集成'],
            ['封装厂商（长电、通富）', '先进封装', '突破集成极限，Chiplet+UCIe'],
            ['云厂商（阿里、腾讯、字节）', '规模化验证', '验证规模化落地，反馈需求'],
            ['研究机构（信通院、清华）', '前沿理论', '提供算法支撑，建立仿真平台']
        ],
        '表3-4 OCS产业协同分工'
    )
    
    add_heading_custom(doc, '（二）培育本土AI芯片+OCS协同生态', 3)
    add_highlight_text(doc, '现状问题：')
    add_body_text(doc, '国内AI芯片生态分散（华为昇腾、壁仞、寒武纪等），缺乏统一互联标准。')
    add_highlight_text(doc, '建议方案：')
    add_body_text(doc, '以OCS为中立互联层，兼容不同AI芯片；参考"光跃超节点"模式（曦智+壁仞+中兴），推广至更多国产AI芯片；制定《国产AI芯片光互联接口规范》团体标准。')
    
    add_heading_custom(doc, '四、标准战略建议', 2)
    
    add_heading_custom(doc, '（一）国内标准布局', 3)
    add_table_custom(doc,
        ['标准层级', '标准名称', '主导单位', '时间节点'],
        [
            ['国家标准', 'OCS光传输拓扑快速认知技术要求', '信通院', '2027年发布'],
            ['行业标准', '数据中心OCS设备技术规范', '开放数据中心委员会', '2026年发布'],
            ['团体标准', '硅光OCS芯片测试方法', '中国通信标准化协会', '2026年发布'],
            ['地方标准', '东数西算枢纽OCS应用指南', '八大枢纽节点', '2027年发布']
        ],
        '表3-5 OCS国内标准布局'
    )
    
    add_heading_custom(doc, '（二）国际标准突破路径', 3)
    add_highlight_text(doc, '目标：')
    add_body_text(doc, '2028年前在OCP OCS项目组中取得实质性话语权。')
    add_highlight_text(doc, '路径：')
    add_body_text(doc, '2026年：积极参与OCP OCS项目组，提交技术白皮书。')
    add_body_text(doc, '2027年：联合光迅、华为等企业，提交拓扑感知等原创技术提案。')
    add_body_text(doc, '2028年：主导1-2项OCS国际标准制定。')
    add_body_text(doc, '2029年：推动中国OCS标准成为国际主流方案之一。')
    
    add_heading_custom(doc, '五、场景应用建议', 2)
    
    add_heading_custom(doc, '（一）基于帕累托前沿的方案选择框架', 3)
    add_highlight_text(doc, '引入SuperPod白皮书帕累托分析框架：')
    add_body_text(doc, '不同OCS架构在"拓扑弹性、故障隔离、功耗与TCO、规模上限、时延、生态成熟度"六个维度上呈现不同的帕累托位置。建议根据具体场景需求，选择最适配的架构。')
    add_table_custom(doc,
        ['维度', '标准构型（电交换）', 'Dragonfly+OCS', '3D Torus+OCS', '大环路+dOCS'],
        [
            ['拓扑弹性', '★', '★★★', '★★★★★', '★★★★★'],
            ['故障隔离', '★★★★★', '★★★★★', '★★★★★', '★★★★★'],
            ['功耗与TCO', '★★★', '★★★★', '★★★★★', '★★★★'],
            ['规模上限', '★★★', '★★★★★', '★★★★', '★★★★'],
            ['时延', '★★★★', '★★★★★', '★★★★', '★★★★'],
            ['生态成熟度', '★★★★★', '★★★★', '★★★', '★★★★']
        ],
        '表3-6 三种OCS构型六维度帕累托对比'
    )
    add_highlight_text(doc, '选择原则：')
    add_body_text(doc, '规模优先（万卡级通用集群）→ Dragonfly+OCS；')
    add_body_text(doc, '能效优先（AI专用集群）→ 3D Torus+OCS；')
    add_body_text(doc, '可靠优先（金融、运营商）→ 大环路+dOCS。')
    
    add_heading_custom(doc, '（二）分场景推进策略', 3)
    add_table_custom(doc,
        ['应用场景', '推荐架构', '成熟度', '推进策略', '目标规模（2030）'],
        [
            ['AI智算中心', '3D Torus+OCS', '高', '优先突破，规模部署', '50%新建智算中心采用OCS'],
            ['东数西算枢纽', '大环路+dOCS', '中', '政策驱动，试点先行', '八大枢纽全覆盖'],
            ['超大规模数据中心', 'Dragonfly+OCS', '中', '头部云厂商引领', '阿里云、腾讯云等规模采购'],
            ['运营商骨干网', '大环路+dOCS', '低', '与OXC协同演进', '重点城市试点'],
            ['6G承载网', '3D Torus+OCS', '低', '前瞻布局，技术储备', '试验网验证']
        ],
        '表3-7 OCS分场景推进策略'
    )
    
    add_heading_custom(doc, '（三）重点工程建议', 3)
    add_highlight_text(doc, '工程一："全光智算"示范工程（2026-2028）')
    add_body_text(doc, '目标：建设10个千卡级全光OCS智算中心示范。投资：预计50亿元。效果：训练效率提升40%，能耗降低40%。')
    add_highlight_text(doc, '工程二："东数西算"全光互联工程（2027-2030）')
    add_body_text(doc, '目标：八大枢纽节点间OCS全光互联。投资：纳入"东数西算"专项，预计100亿元。效果：枢纽间时延小于10ms，带宽提升3倍。')
    add_highlight_text(doc, '工程三：国产AI芯片OCS互联生态工程（2026-2029）')
    add_body_text(doc, '目标：实现5家以上国产AI芯片与OCS协同优化。投资：企业主导，政府补贴30%。效果：打破英伟达NVLink生态垄断。')
    
    # ========== 第四章 ==========
    add_heading_custom(doc, '第四章 政策建议', 1)
    
    add_heading_custom(doc, '一、顶层设计建议', 2)
    
    add_heading_custom(doc, '（一）纳入国家战略', 3)
    add_highlight_text(doc, '建议：')
    add_body_text(doc, '将OCS技术列入《"十五五"国家信息化规划》重大技术方向，与集成电路、人工智能并列。')
    add_highlight_text(doc, '理由：')
    add_body_text(doc, 'OCS是算力基础设施的核心技术；关乎算力安全自主可控；技术窗口期关键，需国家层面统筹。')
    add_highlight_text(doc, '核心判断（引自SuperPod白皮书）：')
    add_quote(doc, '系统能力边界外移速度的差异，会以指数形式放大为系统性能差距。能在制程、封装、互联、精度和软件栈上同时引入新设计变量的参与者，系统算力每两年增长5-6倍；仅在芯片层面优化的参与者，每两年约2-2.5倍。四代之后差距即达16倍。')
    add_body_text(doc, '这意味着OCS不仅是一项技术选择，而是决定中国智算产业能否避免被拉开"难以逆转的差距"的关键变量。')
    
    add_heading_custom(doc, '（二）设立专项基金', 3)
    add_highlight_text(doc, '建议：')
    add_body_text(doc, '设立"全光互联技术与产业创新发展专项"，规模100亿元，分5年实施。')
    add_highlight_text(doc, '重点支持方向：')
    add_body_text(doc, '1. 高端口硅光/MEMS芯片研发（40%）；')
    add_body_text(doc, '2. OCS系统设备产业化（30%）；')
    add_body_text(doc, '3. 应用示范与生态建设（20%）；')
    add_body_text(doc, '4. 标准制定与人才培养（10%）。')
    
    add_heading_custom(doc, '二、产业政策建议', 2)
    
    add_heading_custom(doc, '（一）财税支持', 3)
    add_table_custom(doc,
        ['政策工具', '具体措施', '适用对象'],
        [
            ['研发费用加计扣除', 'OCS芯片研发加计扣除比例提高至150%', '芯片设计企业'],
            ['首台套补贴', '国产OCS设备首台套补贴30%', '设备制造企业'],
            ['示范应用补贴', '采用国产OCS的数据中心补贴PUE改造费用', '数据中心运营方'],
            ['政府采购优先', '政府投资的数据中心优先采用国产OCS', '新建数据中心']
        ],
        '表4-1 OCS产业财税支持政策'
    )
    
    add_heading_custom(doc, '（二）金融支持', 3)
    add_body_text(doc, '产业基金：引导国家集成电路基金、国家制造业转型升级基金投资OCS领域。')
    add_body_text(doc, '信贷支持：对OCS企业提供优惠贷款，利率下浮20%。')
    add_body_text(doc, '上市支持：支持符合条件的OCS企业在科创板、北交所上市。')
    
    add_heading_custom(doc, '三、人才政策建议', 2)
    
    add_heading_custom(doc, '（一）紧缺人才清单', 3)
    add_body_text(doc, '将"硅光芯片设计工程师""OCS系统架构师""光互联网络工程师"列入《国家紧缺人才目录》。')
    
    add_heading_custom(doc, '（二）人才培养计划', 3)
    add_body_text(doc, '高校布局：在10所"双一流"高校设立"硅光技术"微专业。')
    add_body_text(doc, '产教融合：支持光迅、华为等企业与高校共建实验室。')
    add_body_text(doc, '国际引进：对海外OCS领域高端人才给予个税优惠、落户便利。')
    
    # ========== 第五章 ==========
    add_heading_custom(doc, '第五章 风险与应对', 1)
    
    add_heading_custom(doc, '一、主要风险识别', 2)
    add_table_custom(doc,
        ['风险类型', '风险描述', '发生概率', '影响程度'],
        [
            ['技术风险', '大端口芯片技术攻关不及预期', '中', '高'],
            ['市场风险', '电交换技术突破，OCS需求不及预期', '低', '高'],
            ['供应链风险', '高端光刻机、EDA工具断供', '中', '高'],
            ['生态风险', '国产AI芯片生态分散，难以形成合力', '高', '中'],
            ['国际竞争风险', '美国加大技术封锁', '中', '高']
        ],
        '表5-1 OCS产业发展主要风险'
    )
    
    add_heading_custom(doc, '二、应对策略', 2)
    
    add_heading_custom(doc, '（一）技术风险应对', 3)
    add_body_text(doc, '多路线布局：MEMS、硅光、液晶等多技术路线并行，不押注单一路线。')
    add_body_text(doc, '产学研协同：依托国家实验室、高校攻克基础科学问题。')
    add_body_text(doc, '国际合作：与欧洲、日本等非美地区开展技术合作。')
    
    add_heading_custom(doc, '（二）供应链风险应对', 3)
    add_body_text(doc, '国产替代加速：支持上海微电子90nm以下光刻机研发。')
    add_body_text(doc, '工艺创新：发展无需先进制程的硅光技术（90nm即可满足）。')
    add_body_text(doc, '战略储备：建立关键设备、材料战略储备机制。')
    
    add_heading_custom(doc, '（三）生态风险应对', 3)
    add_body_text(doc, '中立互联层：以OCS为中立层，兼容不同AI芯片。')
    add_body_text(doc, '开源社区：建立开源OCS控制软件社区，降低生态门槛。')
    add_body_text(doc, '联盟机制：成立"国产AI芯片光互联联盟"，协同技术路线。')
    
    # ========== 第六章 ==========
    add_heading_custom(doc, '第六章 实施路线图', 1)
    
    add_heading_custom(doc, '一、2026年重点任务', 2)
    add_table_custom(doc,
        ['季度', '重点任务', '责任主体', '交付物'],
        [
            ['Q1', '完成OCS产业发展白皮书', '信通院', '白皮书'],
            ['Q2', '启动"全光智算"示范工程', '工信部', '10个示范项目立项'],
            ['Q3', '发布《数据中心OCS技术规范》', '开放数据中心委员会', '行业标准'],
            ['Q4', '举办首届中国OCS产业大会', '地方政府/行业协会', '产业联盟成立']
        ],
        '表6-1 2026年OCS产业重点任务'
    )
    
    add_heading_custom(doc, '二、2027-2030里程碑', 2)
    add_body_text(doc, '2027年：32×64端口硅光芯片量产；拓扑感知技术成为行业标准；"东数西算"枢纽节点OCS试点部署；国产OCS市场占有率30%。')
    add_body_text(doc, '2028年：64×64/128×128高端口芯片突破；向OCP提交国际标准提案；云厂商、运营商规模采购；国产OCS市场占有率50%。')
    add_body_text(doc, '2029年：OCS+CPO集成方案成熟；主导1-2项OCS国际标准；国产AI芯片+OCS生态形成；国产OCS市场占有率70%。')
    add_body_text(doc, '2030年：256×256+超大端口芯片；全光互联架构普及；中国OCS标准走向国际；全球OCS市场占比35%。')
    
    # ========== 第七章 ==========
    add_heading_custom(doc, '第七章 结语', 1)
    
    add_body_text(doc, 'OCS技术是算力基础设施的"光革命"，是中国实现算力自主可控、换道超车的历史机遇。当前，全球OCS产业正处于从技术验证向规模化商用过渡的关键期，中国具备全产业链优势、巨大市场需求和强有力政策支持，完全有条件在全球OCS产业格局中占据重要地位。')
    
    add_heading_custom(doc, '一、核心判断：系统能力边界外移的竞争', 2)
    add_highlight_text(doc, '引自SuperPod白皮书：')
    add_quote(doc, '超节点竞争的本质不是比拼任何单一规格，而是比拼系统能力边界外移的速度。当前能够独立完成这种全栈联合优化的，只有极少数垂直整合型平台。对多数参与者而言，约束往往不在某项单一技术缺失，而在于芯片、互联、封装、软件、整机与验证尚未形成系统闭环。')
    add_highlight_text(doc, '对中国OCS产业的启示：')
    add_body_text(doc, '不是单点技术追赶：不要试图在MEMS或硅光单条曲线上做线性追赶。')
    add_body_text(doc, '而是系统能力建设：要在芯片、互联、光学、封装、整机、软件与运维的协同工程上建立闭环。')
    add_body_text(doc, '关键窗口期：2026-2028年是技术路线定型关键期，错过将失去话语权。')
    
    add_heading_custom(doc, '二、核心建议总结', 2)
    add_body_text(doc, '1. 技术突破：以"拓扑感知"原创技术为突破口，布局Dragonfly+OCS、3D Torus+OCS、大环路+dOCS三大构型；')
    add_body_text(doc, '2. 标准引领：积极参与OCP OCS项目组，争取2028年前主导1-2项国际标准；')
    add_body_text(doc, '3. 产业协同：构建"国家队+民企"协同创新体，形成从芯片到系统的全栈能力；')
    add_body_text(doc, '4. 场景驱动：以"东数西算"、AI智算中心为抓手，规模拉动产业，实现帕累托最优；')
    add_body_text(doc, '5. 政策支持：设立100亿元专项基金，完善财税、金融、人才政策体系。')
    
    add_heading_custom(doc, '三、三个关键里程碑', 2)
    add_table_custom(doc,
        ['时间节点', '里程碑目标', '标志事件'],
        [
            ['2027年', '并跑发展', '国产OCS市场占有率50%，主导1项国际标准'],
            ['2029年', '生态引领', '国产AI芯片+OCS生态成熟，OCS+CPO集成方案落地'],
            ['2031年', '部分领跑', '256×256+超大端口芯片突破，全球市场份额35%']
        ],
        '表7-1 OCS产业发展关键里程碑'
    )
    
    add_heading_custom(doc, '四、最终判断', 2)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run('时不我待，只争朝夕！')
    set_run_font(run, '黑体', 16, True)
    p.paragraph_format.line_spacing = 1.8
    
    add_body_text(doc, '建议相关部门高度重视OCS技术发展，抓住2026-2028年关键窗口期，统筹规划、加大投入、协同推进。将"拓扑感知"等原创技术上升为国家标准，推动中国OCS产业从"跟跑"到"并跑"再到"部分领跑"，将中国建设成为全球OCS技术创新中心和产业高地，为数字中国、网络强国建设提供坚实支撑。')
    
    # ========== 附录 ==========
    doc.add_page_break()
    add_heading_custom(doc, '附录', 1)
    
    add_heading_custom(doc, '附录A 术语表', 2)
    add_table_custom(doc,
        ['术语', '英文全称', '中文解释'],
        [
            ['OCS', 'Optical Circuit Switching', '光电路交换'],
            ['MEMS', 'Micro-Electro-Mechanical System', '微机电系统'],
            ['CPO', 'Co-Packaged Optics', '共封装光学'],
            ['OXC', 'Optical Cross Connect', '光交叉连接'],
            ['OCP', 'Open Compute Project', '开放计算项目'],
            ['PUE', 'Power Usage Effectiveness', '电能使用效率']
        ],
        '附表1 OCS相关术语表'
    )
    
    add_heading_custom(doc, '附录B 主要参考来源', 2)
    add_body_text(doc, '1. SuperPod白皮书《分布式AI推理加速架构技术白皮书》（DeepLink团队，2025）')
    add_body_text(doc, '   核心贡献：提出系统能力边界外移分析框架、帕累托前沿方法论')
    add_body_text(doc, '   OCS架构：Dragonfly+OCS、3D Torus+OCS、大环路+dOCS三种探索构型')
    add_body_text(doc, '   产业路径：近期/中期/长期三阶段发展策略')
    add_body_text(doc, '2. 中国信通院《中国算力发展指数研究报告》')
    add_body_text(doc, '3. 国家数据局《"数据要素×"三年行动计划》')
    add_body_text(doc, '4. 工信部《算力基础设施高质量发展行动计划》')
    add_body_text(doc, '5. Cignal AI《Optical Circuit Switching Market Report》')
    add_body_text(doc, '6. OCP《OCS Subproject Technical Specification》')
    add_body_text(doc, '7. 各企业公开资料及技术白皮书')
    
    add_heading_custom(doc, '附录C 编制团队', 2)
    add_body_text(doc, '技术顾问：星宿老仙')
    add_body_text(doc, '产业分析：AI助手')
    add_body_text(doc, '数据支持：Tavily搜索平台')
    add_body_text(doc, '文档编制：OpenClaw智能助手')
    
    # 页脚
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_text = '编制完成日期：2026年4月13日\n建议书版本：V1.2.2（政府标准格式版）\n下次修订计划：2026年7月（根据OCP China Day反馈修订）'
    run = footer.add_run(footer_text)
    set_run_font(run, '仿宋_GB2312', 14, False)
    footer.paragraph_format.line_spacing = 1.5
    
    # 保存
    output_path = '/Users/rocky/.openclaw/workspace/reports/中国OCS发展路线建议书_2026_V1.2.2_政府标准版.docx'
    doc.save(output_path)
    print(f"文档已生成：{output_path}")

if __name__ == '__main__':
    main()
