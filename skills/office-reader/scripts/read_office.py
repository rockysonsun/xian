#!/usr/bin/env python3
"""
Office 文档读取工具
支持 .doc, .docx, .xls, .xlsx 文件
"""

import sys
import os

def read_docx(file_path):
    """读取 Word 文档 (.docx)"""
    try:
        from docx import Document
        doc = Document(file_path)
        
        content = []
        # 读取段落
        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text)
        
        # 读取表格
        for table in doc.tables:
            content.append("\n[表格]")
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                content.append(" | ".join(row_text))
            content.append("[表格结束]\n")
        
        return "\n".join(content)
    except ImportError:
        return "错误：未安装 python-docx，请运行：pip install python-docx"
    except Exception as e:
        return f"读取失败：{str(e)}"

def read_doc(file_path):
    """读取旧版 Word 文档 (.doc)"""
    try:
        import subprocess
        # 尝试使用 antiword 或 catdoc
        result = subprocess.run(['antiword', file_path], 
                              capture_output=True, text=True)
        if result.returncode == 0:
            return result.stdout
        
        # 尝试使用 textract
        try:
            import textract
            text = textract.process(file_path)
            return text.decode('utf-8')
        except:
            pass
        
        return "错误：无法读取 .doc 文件，请安装 antiword 或转换为 .docx"
    except Exception as e:
        return f"读取失败：{str(e)}"

def read_xlsx(file_path):
    """读取 Excel 文档 (.xlsx)"""
    try:
        from openpyxl import load_workbook
        wb = load_workbook(file_path, data_only=True)
        
        content = []
        for sheet_name in wb.sheetnames:
            content.append(f"\n=== 工作表: {sheet_name} ===")
            sheet = wb[sheet_name]
            
            for row in sheet.iter_rows():
                row_values = []
                for cell in row:
                    value = cell.value
                    if value is not None:
                        row_values.append(str(value))
                    else:
                        row_values.append("")
                
                # 只输出非空行
                if any(v.strip() for v in row_values):
                    content.append(" | ".join(row_values))
        
        return "\n".join(content)
    except ImportError:
        return "错误：未安装 openpyxl，请运行：pip install openpyxl"
    except Exception as e:
        return f"读取失败：{str(e)}"

def read_xls(file_path):
    """读取旧版 Excel 文档 (.xls)"""
    try:
        import xlrd
        wb = xlrd.open_workbook(file_path)
        
        content = []
        for sheet_idx in range(wb.nsheets):
            sheet = wb.sheet_by_index(sheet_idx)
            content.append(f"\n=== 工作表: {sheet.name} ===")
            
            for row_idx in range(sheet.nrows):
                row_values = []
                for col_idx in range(sheet.ncols):
                    value = sheet.cell_value(row_idx, col_idx)
                    if value:
                        row_values.append(str(value))
                    else:
                        row_values.append("")
                
                if any(v.strip() for v in row_values):
                    content.append(" | ".join(row_values))
        
        return "\n".join(content)
    except ImportError:
        return "错误：未安装 xlrd，请运行：pip install xlrd"
    except Exception as e:
        return f"读取失败：{str(e)}"

def main():
    if len(sys.argv) < 2:
        print("用法: python read_office.py <文件路径>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    
    if not os.path.exists(file_path):
        print(f"错误：文件不存在 {file_path}")
        sys.exit(1)
    
    # 根据扩展名选择读取方式
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.docx':
        print(read_docx(file_path))
    elif ext == '.doc':
        print(read_doc(file_path))
    elif ext == '.xlsx':
        print(read_xlsx(file_path))
    elif ext == '.xls':
        print(read_xls(file_path))
    else:
        print(f"不支持的文件格式：{ext}")
        print("支持的格式：.doc, .docx, .xls, .xlsx")
        sys.exit(1)

if __name__ == '__main__':
    main()
